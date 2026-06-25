#!/usr/bin/env python3
"""Generate deterministic synthetic golden + adversarial fixtures.

Run this once when adding a new format or when a fixture's content needs
to change. Outputs go to `corpus/<format>/` and `corpus/adversarial/<format>/`.
Each fixture is **byte-deterministic** (same script → same SHA-256) so the
MANIFEST.yaml hash check stays meaningful.

Dependencies: `pip install pymupdf python-docx`. These are not pulled by the
spec itself — only maintainers regenerating fixtures need them. CI does NOT
run this script (fixtures are committed).

Usage:
    python clients/extraction/spec/tooling/generate_synthetic_fixtures.py [--format pdf|docx|all]
"""

from __future__ import annotations

import argparse
import io
import sys
import zipfile
from pathlib import Path

SPEC_DIR = Path(__file__).resolve().parent.parent
CORPUS_DIR = SPEC_DIR / "corpus"
ADVERSARIAL_DIR = CORPUS_DIR / "adversarial"

# Determinism: PyMuPDF and python-docx both write timestamps into the
# binary. We override these to make the output reproducible.
FIXED_EPOCH = "20260101000000+00'00'"  # PDF date format
FIXED_ISO = "2026-01-01T00:00:00+00:00"


# ---------------------------------------------------------------------------
# PDF fixtures
# ---------------------------------------------------------------------------


def make_simple_text_pdf(out: Path) -> None:
    """Single-page PDF with one paragraph."""
    import fitz

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_text((72, 100), "Hello world.", fontsize=12)
    page.insert_text((72, 120), "Second line on the same page.", fontsize=12)
    doc.set_metadata(
        {
            "title": "Simple Text PDF",
            "author": "Knovas Synthesizer",
            "creator": "knovas-extract synthetic-fixtures generator",
            "producer": "PyMuPDF",
            "subject": "Golden fixture for text/plain-PDF extraction",
            "keywords": "knovas,extract,fixture,golden",
            "creationDate": f"D:{FIXED_EPOCH}",
            "modDate": f"D:{FIXED_EPOCH}",
        }
    )
    doc.save(str(out), deflate=True, garbage=4, clean=True, deflate_fonts=True)
    doc.close()


def make_multi_page_pdf(out: Path) -> None:
    """Three-page PDF with distinct content per page."""
    import fitz

    doc = fitz.open()
    for i, body in enumerate(
        ["First page text.", "Second page text.", "Third page text."]
    ):
        page = doc.new_page(width=595, height=842)
        page.insert_text((72, 100), f"Page {i + 1}", fontsize=14)
        page.insert_text((72, 130), body, fontsize=11)
    doc.set_metadata(
        {
            "title": "Multi-Page PDF",
            "author": "Knovas Synthesizer",
            "creationDate": f"D:{FIXED_EPOCH}",
            "modDate": f"D:{FIXED_EPOCH}",
        }
    )
    doc.save(str(out), deflate=True, garbage=4, clean=True, deflate_fonts=True)
    doc.close()


def make_encrypted_pdf(out: Path) -> None:
    """Password-protected PDF — extractor must raise EncryptedDocumentError."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), "If you can read this, the test failed.", fontsize=12)
    perm = int(
        fitz.PDF_PERM_ACCESSIBILITY  # allow screen-readers
    )
    doc.save(
        str(out),
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw="owner-pass-do-not-use",
        user_pw="user-pass-do-not-use",
        permissions=perm,
        deflate=True,
    )
    doc.close()


# ---------------------------------------------------------------------------
# DOCX fixtures
# ---------------------------------------------------------------------------


def _strip_docx_timestamps(buf: bytes) -> bytes:
    """python-docx writes a current timestamp into docProps/core.xml; rewrite
    to FIXED_ISO so the binary stays byte-deterministic.

    We re-pack the entire docx with all ZipInfo dates set to a constant so
    even the zip headers don't drift between runs.
    """
    src = zipfile.ZipFile(io.BytesIO(buf), "r")
    out = io.BytesIO()
    dst = zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6)
    try:
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename == "docProps/core.xml":
                # Replace dcterms:created/modified with a fixed timestamp.
                import re

                txt = data.decode("utf-8", errors="replace")
                txt = re.sub(
                    r"<dcterms:created[^>]*>[^<]*</dcterms:created>",
                    f'<dcterms:created xsi:type="dcterms:W3CDTF">{FIXED_ISO}</dcterms:created>',
                    txt,
                )
                txt = re.sub(
                    r"<dcterms:modified[^>]*>[^<]*</dcterms:modified>",
                    f'<dcterms:modified xsi:type="dcterms:W3CDTF">{FIXED_ISO}</dcterms:modified>',
                    txt,
                )
                data = txt.encode("utf-8")
            new_info = zipfile.ZipInfo(info.filename, date_time=(2026, 1, 1, 0, 0, 0))
            new_info.compress_type = zipfile.ZIP_DEFLATED
            dst.writestr(new_info, data)
    finally:
        src.close()
        dst.close()
    return out.getvalue()


def make_simple_docx(out: Path) -> None:
    """Simple DOCX with two paragraphs + metadata."""
    import docx

    d = docx.Document()
    d.add_paragraph("First paragraph of the simple document.")
    d.add_paragraph("Second paragraph with a bit more content.")
    d.core_properties.title = "Simple DOCX"
    d.core_properties.author = "Knovas Synthesizer"
    d.core_properties.subject = "Golden fixture for DOCX extraction"
    d.core_properties.keywords = "knovas,extract,fixture,golden"

    buf = io.BytesIO()
    d.save(buf)
    out.write_bytes(_strip_docx_timestamps(buf.getvalue()))


def make_headings_docx(out: Path) -> None:
    """DOCX with H1 + H2 hierarchy — exercises content.sections[]."""
    import docx

    d = docx.Document()
    d.add_heading("Introduction", level=1)
    d.add_paragraph("Intro paragraph under the H1 heading.")
    d.add_heading("Methods", level=2)
    d.add_paragraph("First methods paragraph.")
    d.add_paragraph("Second methods paragraph.")
    d.add_heading("Results", level=2)
    d.add_paragraph("Results discussion goes here.")
    d.core_properties.title = "Headings DOCX"
    d.core_properties.author = "Knovas Synthesizer"

    buf = io.BytesIO()
    d.save(buf)
    out.write_bytes(_strip_docx_timestamps(buf.getvalue()))


def _info(name: str) -> zipfile.ZipInfo:
    """ZipInfo with deterministic date AND ZIP_DEFLATED compression.

    Without `compress_type = ZIP_DEFLATED`, writestr uses ZIP_STORED regardless
    of the ZipFile's `compression=` kwarg — so a "bomb" payload ends up
    uncompressed and would never trip the decompression-ratio guard.
    """
    z = zipfile.ZipInfo(name, date_time=(2026, 1, 1, 0, 0, 0))
    z.compress_type = zipfile.ZIP_DEFLATED
    return z


def make_decompression_bomb_docx(out: Path) -> None:
    """DOCX with a highly compressible payload triggering the ratio cap.

    Stays small on disk (committed to the repo); expands ~5000× when
    decompressed, well above the default Limits.max_decompression_ratio=100.
    """
    payload = b"A" * (5 * 1024 * 1024)  # 5 MiB of one byte → ~5 KiB compressed
    out_io = io.BytesIO()
    z = zipfile.ZipFile(out_io, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9)
    # Minimal docx scaffold — enough to look like a DOCX to dispatch.
    z.writestr(
        _info("[Content_Types].xml"),
        b'<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"/>',
    )
    z.writestr(
        _info("_rels/.rels"),
        b'<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>',
    )
    z.writestr(_info("word/bomb.bin"), payload)
    z.close()
    out.write_bytes(out_io.getvalue())


def make_zip_slip_docx(out: Path) -> None:
    """DOCX whose central directory carries an entry name with '..' traversal."""
    out_io = io.BytesIO()
    z = zipfile.ZipFile(out_io, "w", compression=zipfile.ZIP_DEFLATED)
    z.writestr(
        _info("[Content_Types].xml"),
        b'<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"/>',
    )
    z.writestr(
        _info("../../../etc/passwd-fake"),
        b"root:x:0:0:fake-passwd-payload\n",
    )
    z.close()
    out.write_bytes(out_io.getvalue())


# ---------------------------------------------------------------------------
# Driver
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# HTML, RTF, EML fixtures (Phase 4)
# ---------------------------------------------------------------------------


def make_simple_html(out: Path) -> None:
    """HTML with title, meta description/keywords, two-level headings."""
    html = (
        '<!DOCTYPE html>\n<html lang="en">\n<head>\n'
        '<meta charset="utf-8">\n<title>Simple HTML</title>\n'
        '<meta name="description" content="Knovas extract golden HTML fixture">\n'
        '<meta name="keywords" content="knovas,extract,fixture,golden">\n'
        "</head>\n<body>\n"
        "<h1>Welcome</h1>\n<p>This is the intro paragraph.</p>\n"
        "<h2>Section A</h2>\n<p>Content of A.</p>\n"
        "<h2>Section B</h2>\n<p>Content of B.</p>\n"
        "</body>\n</html>\n"
    )
    out.write_bytes(html.encode("utf-8"))


def make_js_heavy_html(out: Path) -> None:
    """HTML where most of the body is in a <script> — extractor must drop it."""
    html = (
        '<!DOCTYPE html>\n<html>\n<head><title>JS Heavy</title></head>\n<body>\n'
        "<p>Visible paragraph.</p>\n"
        '<script type="text/javascript">\n'
        'var secrets = "ABCDEFG"; alert("never runs"); '
        'document.write("never runs either");\n'
        "</script>\n"
        "<p>Another visible paragraph.</p>\n"
        "</body>\n</html>\n"
    )
    out.write_bytes(html.encode("utf-8"))


def make_billion_laughs_html(out: Path) -> None:
    """XML-style entity expansion. selectolax doesn't expand entities, so this
    is structurally safe — but we record it as an adversarial fixture so any
    future migration to an XML-strict parser is caught immediately."""
    html = (
        '<?xml version="1.0"?>\n'
        "<!DOCTYPE lolz [\n"
        '  <!ENTITY lol "lol">\n'
        '  <!ENTITY lol2 "&lol;&lol;&lol;&lol;&lol;&lol;&lol;&lol;&lol;&lol;">\n'
        '  <!ENTITY lol3 "&lol2;&lol2;&lol2;&lol2;&lol2;&lol2;&lol2;&lol2;&lol2;&lol2;">\n'
        '  <!ENTITY lol4 "&lol3;&lol3;&lol3;&lol3;&lol3;&lol3;&lol3;&lol3;&lol3;&lol3;">\n'
        "]>\n"
        "<html><body><p>&lol4;</p></body></html>\n"
    )
    out.write_bytes(html.encode("utf-8"))


def make_simple_rtf(out: Path) -> None:
    """Minimal valid RTF with formatted text."""
    rtf = (
        r"{\rtf1\ansi\deff0 "
        r"{\fonttbl {\f0 Times New Roman;}} "
        r"\f0\fs24 Hello \b bold\b0\fs24  RTF world.\par "
        r"Second paragraph with \i italic\i0  text.\par"
        "}"
    )
    out.write_bytes(rtf.encode("ascii"))


def make_object_linking_rtf(out: Path) -> None:
    """RTF with \\object control words (CVE-2017-0199-class regression).

    striprtf has no object loader; the payload bytes are dropped silently.
    Our extractor emits an 'OLE object-linking' warning so callers can audit.
    """
    rtf = (
        r"{\rtf1\ansi\deff0 "
        r"{\object\objemb\objclass{Excel.Sheet}\objw1\objh1 "
        r"{\*\objdata 010500000200000007000000} "
        r"{\result {\rtlch payload-text-here}} } "
        r"\par body text that should still extract.\par"
        "}"
    )
    out.write_bytes(rtf.encode("ascii"))


def make_simple_eml(out: Path) -> None:
    """Plain-text RFC 5322 email."""
    eml = (
        b"From: alice@example.com\r\n"
        b"To: bob@example.com\r\n"
        b"Subject: Test message\r\n"
        b"Date: Thu, 01 Jan 2026 00:00:00 +0000\r\n"
        b"Message-ID: <test-fixture@example.com>\r\n"
        b'Content-Type: text/plain; charset="utf-8"\r\n'
        b"\r\n"
        b"Hello Bob,\r\n\r\n"
        b"This is the body of the test message.\r\n\r\n"
        b"-- \r\n"
        b"Alice\r\n"
    )
    out.write_bytes(eml)


def make_multipart_eml(out: Path) -> None:
    """multipart/alternative with text/plain + text/html — extractor must prefer text/plain."""
    boundary = "knovas-test-boundary"
    eml = (
        f"From: alice@example.com\r\n"
        f"To: bob@example.com\r\n"
        f"Subject: Multipart test\r\n"
        f"Date: Thu, 01 Jan 2026 00:00:00 +0000\r\n"
        f"MIME-Version: 1.0\r\n"
        f'Content-Type: multipart/alternative; boundary="{boundary}"\r\n'
        f"\r\n"
        f"--{boundary}\r\n"
        f'Content-Type: text/plain; charset="utf-8"\r\n\r\n'
        f"Plain version preferred by the extractor.\r\n\r\n"
        f"--{boundary}\r\n"
        f'Content-Type: text/html; charset="utf-8"\r\n\r\n'
        f"<html><body><p>HTML version should be ignored.</p></body></html>\r\n\r\n"
        f"--{boundary}--\r\n"
    ).encode("utf-8")
    out.write_bytes(eml)


# ---------------------------------------------------------------------------
# Driver
# ---------------------------------------------------------------------------


def regenerate(formats: set[str]) -> None:
    if "pdf" in formats:
        (CORPUS_DIR / "pdf").mkdir(exist_ok=True)
        (ADVERSARIAL_DIR / "pdf").mkdir(exist_ok=True)
        make_simple_text_pdf(CORPUS_DIR / "pdf" / "simple-text.pdf")
        make_multi_page_pdf(CORPUS_DIR / "pdf" / "multi-page.pdf")
        make_encrypted_pdf(ADVERSARIAL_DIR / "pdf" / "encrypted.pdf")
        print("PDF: simple-text.pdf, multi-page.pdf, adversarial/encrypted.pdf")

    if "docx" in formats:
        (CORPUS_DIR / "docx").mkdir(exist_ok=True)
        (ADVERSARIAL_DIR / "docx").mkdir(exist_ok=True)
        make_simple_docx(CORPUS_DIR / "docx" / "simple.docx")
        make_headings_docx(CORPUS_DIR / "docx" / "headings.docx")
        make_decompression_bomb_docx(
            ADVERSARIAL_DIR / "docx" / "decompression-bomb.docx"
        )
        make_zip_slip_docx(ADVERSARIAL_DIR / "docx" / "zip-slip.docx")
        print(
            "DOCX: simple.docx, headings.docx, adversarial/decompression-bomb.docx, "
            "adversarial/zip-slip.docx"
        )

    if "html" in formats:
        (CORPUS_DIR / "html").mkdir(exist_ok=True)
        (ADVERSARIAL_DIR / "html").mkdir(exist_ok=True)
        make_simple_html(CORPUS_DIR / "html" / "simple.html")
        make_js_heavy_html(CORPUS_DIR / "html" / "js-heavy.html")
        make_billion_laughs_html(ADVERSARIAL_DIR / "html" / "billion-laughs.html")
        print("HTML: simple.html, js-heavy.html, adversarial/billion-laughs.html")

    if "rtf" in formats:
        (CORPUS_DIR / "rtf").mkdir(exist_ok=True)
        (ADVERSARIAL_DIR / "rtf").mkdir(exist_ok=True)
        make_simple_rtf(CORPUS_DIR / "rtf" / "simple.rtf")
        make_object_linking_rtf(ADVERSARIAL_DIR / "rtf" / "object-linking.rtf")
        print("RTF: simple.rtf, adversarial/object-linking.rtf")

    if "eml" in formats:
        (CORPUS_DIR / "eml").mkdir(exist_ok=True)
        make_simple_eml(CORPUS_DIR / "eml" / "simple.eml")
        make_multipart_eml(CORPUS_DIR / "eml" / "multipart-alternative.eml")
        print("EML: simple.eml, multipart-alternative.eml")


def main(argv: list[str] | None = None) -> int:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument(
        "--format",
        choices=("pdf", "docx", "html", "rtf", "eml", "all"),
        default="all",
        help="Which family to regenerate.",
    )
    args = p.parse_args(argv)
    formats = (
        {"pdf", "docx", "html", "rtf", "eml"}
        if args.format == "all"
        else {args.format}
    )
    try:
        regenerate(formats)
    except ImportError as exc:
        print(
            f"ERROR: missing generator dep ({exc.name}). "
            "Install with: pip install pymupdf python-docx",
            file=sys.stderr,
        )
        return 2
    print(
        "\nNext: run `python clients/extraction/spec/tooling/hash_fixtures.py` "
        "to refresh MANIFEST.yaml hashes."
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
